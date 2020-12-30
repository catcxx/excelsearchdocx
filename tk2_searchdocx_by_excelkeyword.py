#!/usr/bin/env python
# -*- coding: utf-8 -*-

import os
import sys
import time
import xlrd
import xlwt
import docx
import re
from tkinter import *
from win32com import client as wc

LOG_LINE_NUM = 0
class MY_GUI():
    def __init__(self,init_window_name):
        self.init_window_name = init_window_name

    #设置窗口
    def set_init_window(self):
        self.init_window_name.title("excel&word关键字搜索工具_v1.0")           #窗口名
        #self.init_window_name.geometry('320x160+10+10')                         #290 160为窗口大小，+10 +10 定义窗口弹出时的默认展示位置
        self.init_window_name.geometry('1068x681+10+10')
        #self.init_window_name["bg"] = "pink"                                    #窗口背景色，其他背景色见：blog.csdn.net/chl0000/article/details/7657887
        #self.init_window_name.attributes("-alpha",0.9)                          #虚化，值越小虚化程度越高
        #标签
        self.init_data_label_docx = Label(self.init_window_name, text="待传入docxdir")
        self.init_data_label_docx.grid(row=0, column=0)
        self.init_data_label_excel = Label(self.init_window_name, text="待传入excelfilepath")
        self.init_data_label_excel.grid(row=0, column=10)
        # self.result_data_label = Label(self.init_window_name, text="输出结果")
        # self.result_data_label.grid(row=0, column=12)
        self.log_label = Label(self.init_window_name, text="日志")
        self.log_label.grid(row=12, column=0)
        #文本框
        self.init_data_Text_docx = Text(self.init_window_name, width=67, height=15)  #原始数据录入框
        self.init_data_Text_docx.grid(row=1, column=0, rowspan=10, columnspan=10)
        self.init_data_Text_excel = Text(self.init_window_name, width=67, height=15)  #原始数据录入框
        self.init_data_Text_excel.grid(row=1, column=10, rowspan=10, columnspan=10)

        # self.result_data_Text = Text(self.init_window_name, width=70, height=49)  #处理结果展示
        # self.result_data_Text.grid(row=1, column=12, rowspan=15, columnspan=10)
        self.log_data_Text = Text(self.init_window_name, width=66, height=9)  # 日志框
        self.log_data_Text.grid(row=13, column=0, columnspan=10)
        #按钮

        # print()
        self.excel_worddir_search_button = Button(self.init_window_name, text="搜索", bg="lightblue", width=10,command=self.excel_worddir_search)  # 调用内部方法  加()为直接调用
        self.excel_worddir_search_button.grid(row=11, column=12)


    #功能函数
    #从excel文件row[2] 的第二行开始读取关键字
    def read_keyword_from_excel(self):
        #excelpath = self.init_data_Text_excel.get(1.0,END)
        # print("----in read_keyword_from_excel---")
        # print(self.init_data_Text_excel.get(1.0,END))
        # 打开excel文件，创建一个workbook对象,book对象也就是fruits.xlsx文件,表含有sheet名
        # path = os.path.j
        rbook = xlrd.open_workbook(self.init_data_Text_excel.get(1.0,END).strip().replace("\n","").encode())
        # sheets方法返回对象列表,[<xlrd.sheet.Sheet object at 0x103f147f0>]
        rbook.sheets()
        # xls默认有3个工作簿,Sheet1,Sheet2,Sheet3
        rsheet = rbook.sheet_by_index(0)  # 取第一个工作簿

        keyword_list = []
        

        # 循环工作簿的所有行
        for row in rsheet.get_rows():
            keyword1_column = row[1]  # 所在的列
            keyword1_value = keyword1_column.value  #
            if keyword1_value != '单位':  # 排除第一行
                # price_column = row[2]  # 价格所在的列
                # price_value = price_column.value
                # 打印
                # print("模块", keyword1_value, "价格", price_value)
                keyword_list.append(keyword1_value)

        print(keyword_list)
        return keyword_list

    ###################




    def read_word(self):
        document = docx.Document(docxfile)
        pattern = re.compile(self)
        print(pattern)
        for paragraph in document.paragraphs:
            if re.search(pattern, paragraph.text):
                return True
        for table in document.tables:
            for row_idx in range(0, len(table.rows)):
                for cell in table.row_cells(row_idx):
                    if re.search(pattern, cell.text):
                        return True
        return False
    ############################################

    def getfilelist(self):
        #docxfilepath = self.init_data_Text_docx.get(1.0,END).strip().replace("\n","").encode()
        filelist =  os.listdir(self.init_data_Text_docx.get(1.0,END).strip().replace("\n","").encode("UTF-8")) 
        print(filelist)
        files = []

        word = wc.Dispatch("Word.Application")
        for i in range(len(filelist)):
            if  filelist[i].decode().endswith(".doc"):
                print(filelist[i].decode())
                child = os.path.join('%s\\%s' % (self.init_data_Text_docx.get(1.0,END).strip().replace("\n",""), filelist[i].decode()))
                # print(child)
                doc = word.Documents.Open(child)#打开word文件
                doc.SaveAs("{}x".format(child), 12)#另存为后缀为".docx"的文件，其中参数12指docx文件
                doc.Close() #关闭原来word文件
        word.Quit()

        filelist =  os.listdir(self.init_data_Text_docx.get(1.0,END).strip().replace("\n","").encode("UTF-8")) 
        for i in range(len(filelist)):
            print("\\\\"+filelist[i].decode())
            print(self.init_data_Text_docx.get(1.0,END).strip().replace("\n",""))
            child = os.path.join('%s\\%s' % (self.init_data_Text_docx.get(1.0,END).strip().replace("\n",""), filelist[i].decode()))

            if filelist[i].decode().endswith(".docx"):
                if os.path.isdir(child):
                    files.extend(getfilelist(child))
                else:
                    files.append(child)
        print(files)
        
        return files

    #docxfilepath:docx dir, egg:r'C:\Users\ch\Desktop\test\docx';
    #excelpath:excel filepath, egg: r'C:\Users\ch\Desktop\test\1.xls'
    #从excel列读取关键字， 到word文件目录搜索，把结果写入文件
    def excel_worddir_search(self):
        docxfilepath = self.init_data_Text_docx.get(1.0,END).strip().replace("\n","").encode()
        excelpath = self.init_data_Text_excel.get(1.0,END).strip().replace("\n","").encode()
        result_list = []
        # print("in excel_worddir_search-------------------hhh ")
        # print(self.init_data_Text_docx.get(1.0,END))
        # print(self.init_data_Text_excel.get(1.0,END))
        if (docxfilepath):
            try:
                keyword_list = self.read_keyword_from_excel()
                print(keyword_list)
                docxfile_list = self.getfilelist()
                print("docxfile_list")
                counter = 0
                for file in docxfile_list:
                    document = docx.Document(file)
                    counter +=1
                    print("----finished---"+'percent: {:.2%}'.format(counter/len(docxfile_list)))
                    self.write_log_to_Text("INFO:"+'percent: {:.2%}'.format(counter/len(docxfile_list)))
                    text = ''
                    for paragraph in document.paragraphs:
                        text += paragraph.text
                    for table in document.tables:
                        for row_idx in range(0, len(table.rows)):
                            for cell in table.row_cells(row_idx):
                                text += cell.text
                    for keyword in keyword_list:     
                        res = 0
                        pattern = re.compile(keyword)
                        # print(pattern)
                        if re.search(pattern, text):
                            res = 1
                        if res:
                            print("keyword："+keyword +" appears in：" + file)
                            result_list.append(keyword+"-"+file)                    
            except:
                print("error：")

        print(sorted(result_list))
        #write result into txt
        with open('result.txt', 'w',encoding= 'utf-8') as f:
            for item in sorted(result_list):
                f.write(item+"\r")



    def get_current_time(self):
        current_time = time.strftime('%Y-%m-%d %H:%M:%S',time.localtime(time.time()))
        return current_time

    #write log into loglabel
    def write_log_to_Text(self,logmsg):
        global LOG_LINE_NUM
        current_time = self.get_current_time()
        logmsg_in = str(current_time) +" " + str(logmsg) + "\n"
        if LOG_LINE_NUM <= 7:
            self.log_data_Text.insert(END, logmsg_in)
            LOG_LINE_NUM = LOG_LINE_NUM + 1
        else:
            self.log_data_Text.delete(1.0,2.0)
            self.log_data_Text.insert(END, logmsg_in)

def gui_start():
    init_window = Tk()              
    ZMJ_PORTAL = MY_GUI(init_window)
    ZMJ_PORTAL.set_init_window()

    init_window.mainloop()


gui_start()